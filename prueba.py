# from passlib.context import CryptContext

# pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")
# hashed = pwd_context.hash("root")
# print(hashed)

import requests
import os
from enum import Enum
from app.schemas.players import PlayerCreate
from docx import Document
import requests
from requests.auth import HTTPBasicAuth
from io import BytesIO
from docx import Document
import re
import xml.etree.ElementTree as ET
from typing import List
from app.utils.student_verification import sigenu_check


def fetch_players_from_nextcloud() -> List[PlayerCreate]:
    USUARIO = os.getenv("USER_NEXTCLOUD") or "kevin.ortega@rect.uh.cu"
    CONTRASEÑA_APP = os.getenv("PASS_NEXTCLOUD") or "1PassdeKevin."

    WEBDAV_URL = f"https://minube.uh.cu/remote.php/dav/files/{USUARIO}/caribes.uh.cu/"

    r = requests.request(
        "PROPFIND",
        WEBDAV_URL,
        auth=HTTPBasicAuth(USUARIO, CONTRASEÑA_APP),
        headers={"Depth": "1"},
        timeout=20
    )
    if r.status_code != 207:
        raise Exception(f"Error al listar archivos: {r.status_code}", [])
    

    ns = {"d": "DAV:"}
    root = ET.fromstring(r.content)
    docx_files = []
    for resp in root.findall("d:response", ns):
        href_elem = resp.find("d:href", ns)
        if href_elem is not None and href_elem.text:
            href = href_elem.text
            if href.endswith(".docx") and "/." not in href:
                filename = href.split("/")[-1]
                if filename:
                    docx_files.append(filename)

    all_players: List[PlayerCreate] = []

    for filename in sorted(docx_files):
        file_url = f"{WEBDAV_URL}{filename}"
        r = requests.get(file_url, auth=HTTPBasicAuth(USUARIO, CONTRASEÑA_APP), timeout=15)
        r.raise_for_status()
        doc = Document(BytesIO(r.content))

        # Extraer facultad
        faculty = None
        for p in doc.paragraphs[:10]:
            text = p.text.replace("F acultad", "Facultad").strip()
            if "Facultad:" in text:
                m = re.search(r"Facultad:\s*([^\t\n\r]+?)(?:\s*Deporte:|\t|$)", text)
                if m:
                    faculty = m.group(1).strip()
                break

        if not doc.tables:
            continue
        table = doc.tables[0]
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if len(cells) < 2:
                continue
            name, ci = cells[0], cells[1]
            if i == 0 or not name or len(name) < 3 or not any(c.isdigit() for c in ci):
                continue
            all_players.append(PlayerCreate(CI=ci, name=name, faculty=faculty))

    return all_players

players = fetch_players_from_nextcloud()
repeat = []
no =[]
i=0
for player in players:
    if player.CI in repeat:
        continue
    if not sigenu_check(player.CI):
        no.append(player)
        continue
    print(i)
    print(player)
    repeat.append(player.CI)
    i+=1

print(no)