import requests
import os
from enum import Enum
from app.schemas.players import PlayerCreate
from docx import Document
import requests
from requests.auth import HTTPBasicAuth
from io import BytesIO
from docx import Document
import re
import xml.etree.ElementTree as ET
from typing import List

class SportEnum(Enum):
    KICKINGBALL = 0
    HANDBALL = 1
    TRACCION_SOGA = 2
    TRAGABOLAS = 3
    LANZAMIENTO_DARDOS = 4
    DOMINO = 5
    CARRERA_SACOS = 6
    ANILLAS = 7
    JUEGOS_MIXTOS = 8
    FUTBOL = 9
    BALONCESTO = 10
    VOLEIBOL = 11
    
    @classmethod
    def get_name_by_number(cls, number: int):
        """Obtener el nombre formateado (más legible)"""
        mapping = {
            0: "Kickingball femenino",
            1: "Handball masculino",
            2: "Tracción de la soga",
            3: "Tragabolas",
            4: "Lanzamiento de dardos",
            5: "Dominó",
            6: "Carrera de sacos",
            7: "Anillas",
            8: "Juegos mixtos",
            9: "Fútbol",
            10: "Baloncesto", 
            11: "Voleibol"
        }
        return mapping.get(number, f"Desconocido ({number})")

class FacultadEnum(Enum):
    """Enum para mapear IDs numéricos a nombres de facultades"""
    
    # Definición de todas las facultades con IDs consecutivos
    MATEMATICA_COMPUTACION = 0
    DERECHO = 1
    BIOLOGIA = 2
    LENGUAS_EXTRANJERAS = 3
    FISICA = 4
    FARMACIA_ALIMENTOS = 5
    FILOSOFIA_HISTORIA_SOCIOLOGIA = 6
    ECONOMIA = 7
    GEOGRAFIA = 8
    DISENO = 9
    COMUNICACION_SOCIAL = 10
    ARTES_LETRAS = 11
    QUIMICA = 12
    CONTABILIDAD = 13
    PSICOLOGIA = 14
    TURISMO = 15
    
    @classmethod
    def obtener_nombre_formateado(cls, numero: int) -> str:
        """
        Obtener el nombre completo y formateado para mostrar
        
        Args:
            numero: ID numérico de la facultad (1-16)
            
        Returns:
            str: Nombre completo formateado
        """
        mapping = {
            0: "Facultad de Matemática y Computación",
            1: "Facultad de Derecho",
            2: "Facultad de Biología",
            3: "Facultad de Lenguas Extranjeras",
            4: "Facultad de Física",
            5: "Instituto de Farmacia y Alimentos",
            6: "Facultad de Filosofía, Historia y Sociología",
            7: "Facultad de Economía",
            8: "Facultad de Geografía",
            9: "Instituto de Diseño",
            10: "Facultad de Comunicación Social",
            11: "Facultad de Artes y Letras",
            12: "Facultad de Química",
            13: "Facultad de Contabilidad",
            14: "Facultad de Psicología",
            15: "Facultad de Turismo"
        }
        return mapping.get(numero, f"Facultad Desconocida ({numero})")


def extract_players():
    user = os.getenv("USER_NEXTCLOUD")
    password= os.getenv("PASS_NEXTCLOUD")
    url = "https://minube.uh.cu/apps/tables/api/1/tables/263/rows"
    
    auth = (user, password) if user and password else None

    response = requests.get(url, auth=auth, headers={"Accept": "application/json"})
    response.raise_for_status()

    raw_rows = response.json()

    rows = []
    for row in raw_rows:
        fila = {col["columnId"]: col["value"] for col in row["data"]}
        rows.append(fila)
    return rows



def fetch_players_from_nextcloud() -> List[PlayerCreate]:
    USUARIO = os.getenv("USER_NEXTCLOUD") or ""
    CONTRASEÑA_APP = os.getenv("PASS_NEXTCLOUD") or ""

    WEBDAV_URL = f"https://minube.uh.cu/remote.php/dav/files/{USUARIO}/caribes.uh.cu/"

    r = requests.request(
        "PROPFIND",
        WEBDAV_URL,
        auth=HTTPBasicAuth(USUARIO, CONTRASEÑA_APP),
        headers={"Depth": "1"},
        timeout=20
    )
    if r.status_code != 207:
        raise Exception(f"Error al listar archivos: {r.status_code}", [])
    

    ns = {"d": "DAV:"}
    root = ET.fromstring(r.content)
    docx_files = []
    for resp in root.findall("d:response", ns):
        href_elem = resp.find("d:href", ns)
        if href_elem is not None and href_elem.text:
            href = href_elem.text
            if href.endswith(".docx") and "/." not in href:
                filename = href.split("/")[-1]
                if filename:
                    docx_files.append(filename)

    all_players: List[PlayerCreate] = []

    for filename in sorted(docx_files):
        file_url = f"{WEBDAV_URL}{filename}"
        r = requests.get(file_url, auth=HTTPBasicAuth(USUARIO, CONTRASEÑA_APP), timeout=15)
        r.raise_for_status()
        doc = Document(BytesIO(r.content))

        # Extraer facultad
        faculty = None
        for p in doc.paragraphs[:10]:
            text = p.text.replace("F acultad", "Facultad").strip()
            if "Facultad:" in text:
                m = re.search(r"Facultad:\s*([^\t\n\r]+?)(?:\s*Deporte:|\t|$)", text)
                if m:
                    faculty = m.group(1).strip()
                break

        if not doc.tables:
            continue
        table = doc.tables[0]
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if len(cells) < 2:
                continue
            name, ci = cells[0], cells[1]
            if i == 0 or not name or len(name) < 3 or not any(c.isdigit() for c in ci):
                continue
            all_players.append(PlayerCreate(CI=ci, name=name, faculty=faculty))

    return all_players